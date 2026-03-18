import os, io, re, logging
from pathlib import Path
import fitz
import anthropic
import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, ContextTypes, filters

logging.basicConfig(format="%(asctime)s | %(levelname)s | %(message)s", level=logging.INFO)
logger = logging.getLogger(__name__)

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "8625895538:AAGc4xYBK4J0t0yq9KISQsE8QFfJR10StMo")
ANTHROPIC_API_KEY  = os.getenv("ANTHROPIC_API_KEY",  "YOUR_ANTHROPIC_API_KEY")

claude = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

SUMMARY_MODES = {
    "short":    {"ar": "ملخص قصير (5-7 نقاط)"},
    "detailed": {"ar": "ملخص مفصّل مع عناوين"},
    "bullets":  {"ar": "نقاط رئيسية فقط"},
    "academic": {"ar": "ملخص أكاديمي رسمي"},
}

def extract_text_from_pdf(file_bytes):
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()

def extract_text_from_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def extract_text_from_url(url):
    try:
        resp = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        text = re.sub(r"<[^>]+>", " ", resp.text)
        return re.sub(r"\s+", " ", text).strip()[:15000]
    except Exception as e:
        return f"[تعذّر جلب الرابط: {e}]"

def summarize_text(text, mode="detailed"):
    mode_label = SUMMARY_MODES.get(mode, SUMMARY_MODES["detailed"])["ar"]
    response = claude.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=2000,
        system="أنت مساعد متخصص في تلخيص المحتوى باللغة العربية. ردودك دائماً بالعربية الفصيحة مع عناوين ونقاط منظمة.",
        messages=[{"role": "user", "content": f"قم بعمل {mode_label} للنص التالي:\n\n{text[:12000]}"}],
    )
    return response.content[0].text

def create_summary_docx(summary, source_title, mode):
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("📄 ملخص بالذكاء الاصطناعي")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    doc.add_paragraph("─" * 60)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = info.add_run("📌 المصدر: ")
    r.bold = True
    info.add_run(source_title)

    doc.add_paragraph("─" * 60)

    for line in summary.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if line.startswith("# "):
            r = para.add_run(line[2:])
            r.bold = True; r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        elif line.startswith("## "):
            r = para.add_run(line[3:])
            r.bold = True; r.font.size = Pt(13)
        elif line.startswith(("- ", "• ", "* ")):
            para.add_run("• " + line[2:]).font.size = Pt(11)
        else:
            para.add_run(line).font.size = Pt(11)

    doc.add_paragraph("─" * 60)
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run("تم إنشاؤه بواسطة بوت التلخيص الذكي 🤖")
    fr.font.size = Pt(9)
    fr.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

async def cmd_start(update, context):
    await update.message.reply_text(
        "👋 *أهلاً بك في بوت التلخيص الذكي!*\n\n"
        "أرسل لي:\n• ✍️ نصاً\n• 📄 ملف PDF\n• 📝 ملف Word\n• 🔗 رابط موقع\n\n"
        "وسأرسل لك ملخصاً كملف Word! 📎\n\n"
        "/mode - اختيار نمط التلخيص",
        parse_mode="Markdown"
    )

async def cmd_mode(update, context):
    keyboard = [
        [InlineKeyboardButton("📋 ملخص قصير",    callback_data="mode_short")],
        [InlineKeyboardButton("📖 ملخص مفصّل",   callback_data="mode_detailed")],
        [InlineKeyboardButton("🔵 نقاط رئيسية",  callback_data="mode_bullets")],
        [InlineKeyboardButton("🎓 ملخص أكاديمي", callback_data="mode_academic")],
    ]
    await update.message.reply_text("📝 اختر نمط التلخيص:", reply_markup=InlineKeyboardMarkup(keyboard))

async def callback_handler(update, context):
    query = update.callback_query
    await query.answer()
    if query.data.startswith("mode_"):
        mode = query.data[5:]
        context.user_data["mode"] = mode
        await query.edit_message_text(f"✅ تم اختيار: *{SUMMARY_MODES[mode]['ar']}*", parse_mode="Markdown")

async def handle_text(update, context):
    text = update.message.text.strip()
    if re.match(r"https?://\S+", text):
        await update.message.reply_text("🔗 جاري جلب الرابط...")
        extracted = extract_text_from_url(text)
        title = text[:50]
    else:
        if len(text) < 50:
            await update.message.reply_text("⚠️ النص قصير جداً!")
            return
        extracted, title = text, text[:50]
    await _do_summarize(update.message, context, extracted, title)

async def handle_document(update, context):
    doc = update.message.document
    ext = Path(doc.file_name or "").suffix.lower()
    if ext not in (".pdf", ".docx", ".txt"):
        await update.message.reply_text("⚠️ يُقبل فقط: PDF, DOCX, TXT")
        return
    await update.message.reply_text("📥 جاري تحميل الملف...")
    file = await context.bot.get_file(doc.file_id)
    file_bytes = bytes(await file.download_as_bytearray())
    try:
        if ext == ".pdf":    extracted = extract_text_from_pdf(file_bytes)
        elif ext == ".docx": extracted = extract_text_from_docx(file_bytes)
        else:                extracted = file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        await update.message.reply_text(f"❌ خطأ: {e}")
        return
    await _do_summarize(update.message, context, extracted, doc.file_name)

async def _do_summarize(message, context, text, title):
    mode = context.user_data.get("mode", "detailed")
    await message.reply_text("⏳ جاري التلخيص بالذكاء الاصطناعي...")
    try:
        summary = summarize_text(text, mode=mode)
        docx_bytes = create_summary_docx(summary, title, mode)
        safe = re.sub(r"[^\w\u0600-\u06FF]", "_", title)[:30]
        await message.reply_document(
            document=io.BytesIO(docx_bytes),
            filename=f"ملخص_{safe}.docx",
            caption=f"✅ *تم التلخيص!*\nالنمط: {SUMMARY_MODES[mode]['ar']}",
            parse_mode="Markdown"
        )
        preview = summary[:600] + "\n\n📄 *[الملخص الكامل في الملف أعلاه]*"
        await message.reply_text(f"📝 *معاينة:*\n\n{preview}", parse_mode="Markdown")
    except Exception as e:
        await message.reply_text(f"❌ حدث خطأ: {e}")

def main():
    app = Application.builder().token(TELEGRAM_BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CommandHandler("mode",  cmd_mode))
    app.add_handler(CallbackQueryHandler(callback_handler))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    print("✅ البوت يعمل!")
    app.run_polling(allowed_updates=Update.ALL_TYPES)

if __name__ == "__main__":
    main()
